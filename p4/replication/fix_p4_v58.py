"""
P4 Vietnam v5.8 → v5.9 patches
================================
Apply 7 post-revision fixes to BOTH manuscript_blinded.docx and
manuscript_full_with_authors.docx in the IJoEM submission package.

Issues identified during post-v5.8 audit (all carried over from earlier
revisions, never previously fixed):

1. **Metadata Tables/Figures count mismatch** — header line claims
   "Tables: 2 ... Figures: 2" but body actually has 4 tables (descriptives,
   directional summary by wave, robustness, turning points) and 3 figures
   (conceptual model, predicted I–P curves with 4 panels, moderation slices).

2. **§4.4 duplicate "Table 1" label** — body has both a Table 1 (descriptives,
   in §4.0) AND a "Table 1" referenced in §4.4 directional summary. The
   second one should be Table 2.

3. **§4.5 ad-hoc "Table LM" name** — should be Table 4 to match consecutive
   numbering (Table 1 descriptives, 2 directional, 3 robustness, 4 turning
   points).

4. **Supplementary materials list** — references CSV by old "Table LM source"
   label; updated to "Table 4 source" (CSV filename itself unchanged).

5. **§2.1 missing Cuervo-Cazurra (2012) cite** — reference is in the
   bibliography but never cited in body. Adding it where Vietnam EMNE setup
   is discussed restores citation consistency.

6. **§2.4 missing Helfat & Peteraf (2003) cite** — reference is in the
   bibliography but never cited in body. The dynamic-capability lifecycle
   perspective fits naturally into the stage-contingent digital value
   argument in §2.4.

7. **Title discoverability** — append "— Three-Wave Evidence from Vietnam"
   so search engines and reviewers can find the country setting from the
   title alone.

Note: P1 highlights count (5 vs 6) and P2 inline stats consistency were
checked in audit and turned out to be already correct (5 highlights, all
β/p stats spaced consistently per APA 7).

Usage:
    pip install python-docx
    python3 fix_p4_v58.py

Outputs:
    manuscript_blinded_FIXED.docx
    manuscript_full_with_authors_FIXED.docx
"""
import shutil
from docx import Document


def replace_in_para(p, old, new):
    """Replace text in paragraph; preserves run-level formatting where possible."""
    if old not in p.text:
        return False
    for run in p.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    new_text = p.text.replace(old, new)
    for run in p.runs[1:]:
        run.text = ""
    if p.runs:
        p.runs[0].text = new_text
    return True


def apply_fixes(src, dst):
    """Apply all 7 fixes to a manuscript file."""
    shutil.copy(src, dst)
    doc = Document(dst)
    changes = []

    REPLACEMENTS = [
        # 1. Metadata Tables/Figures count
        ('Tables: 2 (Table 1 descriptives in supplementary appendix; Table 2 main empirical pattern in §4.4). Figures: 2 (conceptual model and predicted I–P curves).',
         'Tables: 4 (Table 1 analytic-sample descriptives in §4.0; Table 2 directional interpretation of focal coefficients by wave and pooled in §4.4; Table 3 robustness collation in §4.5; Table 4 implied turning points and Lind–Mehlum p-values in §4.5). Figures: 3 (Figure 1 conceptual model; Figure 2 with four panels showing predicted I–P curves for 2009 / 2015 / 2023 / pooled; Figure 3 moderation slices at p25 vs p75 of TCI_z and DAI_z).'),

        # 2. §4.4 duplicate Table 1 → Table 2
        ('Table 1 summarises the directional interpretation of the focal coefficients',
         'Table 2 summarises the directional interpretation of the focal coefficients'),

        # 3. §4.5 Table LM → Table 4
        ('Table LM reports the implied turning points',
         'Table 4 reports the implied turning points'),

        # 4. Supplementary materials CSV ref
        ('tables/table_lind_mehlum.csv (Table LM source)',
         'tables/table_lind_mehlum.csv (Table 4 source)'),

        # 5. §2.1 add Cuervo-Cazurra (2012)
        ('transitional economy such as Vietnam',
         'transitional economy such as Vietnam (Cuervo-Cazurra, 2012)'),

        # 6. §2.4 add Helfat & Peteraf (2003)
        ('In early phases of internationalisation, digital tools may create relatively direct gains',
         'Drawing on the dynamic-capability lifecycle perspective (Helfat & Peteraf, 2003), we argue that in early phases of internationalisation, digital tools may create relatively direct gains'),
    ]

    # Title (separate handling for exact match on first paragraph)
    TITLE_OLD = 'Revisiting the Internationalisation–Performance Relationship in an Emerging Market: The Roles of Technological Capability and Digital Adoption'
    TITLE_NEW = TITLE_OLD + ' — Three-Wave Evidence from Vietnam'

    for i, p in enumerate(doc.paragraphs):
        for old, new in REPLACEMENTS:
            if old in p.text:
                if replace_in_para(p, old, new):
                    changes.append(f'  Para {i}: {old[:50]}...')

    # Title fix
    for i, p in enumerate(doc.paragraphs[:5]):
        if p.text.strip() == TITLE_OLD:
            if len(p.runs) == 1:
                p.runs[0].text = TITLE_NEW
            else:
                for run in p.runs[1:]:
                    run.text = ""
                p.runs[0].text = TITLE_NEW
            changes.append(f'  Para {i}: TITLE updated')
            break

    # Tables/footers also need check
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for old, new in REPLACEMENTS:
                        if old in p.text:
                            if replace_in_para(p, old, new):
                                changes.append(f'  Table cell: {old[:50]}...')

    doc.save(dst)
    return changes


if __name__ == '__main__':
    for kind in ['blinded', 'full_with_authors']:
        src = f'submission/manuscript_{kind}.docx'
        dst = f'submission/manuscript_{kind}_FIXED.docx'
        print(f'=== Fixing {kind} ===')
        changes = apply_fixes(src, dst)
        for c in changes:
            print(c)
        print(f'  -> {dst}')
        print(f'  Total: {len(changes)} edits\n')
