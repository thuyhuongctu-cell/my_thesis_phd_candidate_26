import os
import re
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import seaborn as sns
import nltk
import torch

# ========== TRANSFORMERS & RAKE ========== #
from transformers import pipeline
from rake_nltk import Rake

# ========== Sentence Embeddings ========== #
from sentence_transformers import SentenceTransformer, util as st_utils

# ========== Word & Doc Handling ========== #
from docx import Document
from wordcloud import WordCloud

# ========== Para barras de progreso ========== #
from tqdm import tqdm

# ========== Para redes de colaboración ========== #
import networkx as nx
import plotly.express as px

import unicodedata

# =========================
#        CONFIG
# =========================
SCOPUS_PATH = r"D:\El investigador\46_GOLD_PREG_ROBB\scopus.csv"
WOS_PATH    = r"D:\El investigador\46_GOLD_PREG_ROBB\wos.xlsx"
OUTPUT_DIR  = r"D:\El investigador\46_GOLD_PREG_ROBB\Outputs21"

os.makedirs(OUTPUT_DIR, exist_ok=True)

# ========== MODELO DE EMBEDDINGS ========== #
embedding_model = SentenceTransformer('all-MiniLM-L6-v2')

# Verificamos si hay GPU disponible
device_id = 0 if torch.cuda.is_available() else -1

# ========== SUMMARIZATION PIPELINE ========== #
#
# IMPORTANTE: El mensaje:
# "Some weights of PegasusForConditionalGeneration were not initialized..."
# es solo una advertencia de HuggingFace al cargar el modelo. No es un error 
# y el código seguirá ejecutándose correctamente.
# 
# Si deseas suprimirlo, se pueden usar argumentos adicionales 
# en pipeline, pero no es necesario para que funcione.
#
summarizer = pipeline(
    "summarization",
    model="google/pegasus-xsum",
    device=device_id  
)

# ==========================================================
# TÉRMINOS BASE + OBTENCIÓN DE SINÓNIMOS (WordNet)
# ==========================================================
BASE_TERMS = [
    "gold", "pretreat", "roast", "dielectric heating", "microwave-assisted", "preg robb", "double refractor", "carbonaceous matter", "graphit", "kerogen", "Carlin-type",
]

from nltk.corpus import stopwords, wordnet
from nltk.tokenize import word_tokenize
from sklearn.cluster import KMeans
from sklearn.metrics import silhouette_score
from sklearn.feature_extraction.text import CountVectorizer

def get_wordnet_synonyms(term, max_syn=3):
    synonyms = []
    tokens = term.split()
    # Si hay más de una palabra, evitamos buscar sinónimos en WordNet
    if len(tokens) > 1:
        return synonyms
    
    for syn in wordnet.synsets(term.lower()):
        for lemma in syn.lemmas():
            candidate = lemma.name().replace("_"," ")
            if candidate.lower() != term.lower() and candidate not in synonyms:
                synonyms.append(candidate)
                if len(synonyms) >= max_syn:
                    break
        if len(synonyms) >= max_syn:
            break
    return synonyms

def build_term_weights(base_terms):
    term_weights = {}
    for term in base_terms:
        term_weights[term.lower()] = 2.0
        syns = get_wordnet_synonyms(term)
        for s in syns:
            if s.lower() not in term_weights:
                term_weights[s.lower()] = 2.0
    return term_weights

TERM_WEIGHTS = build_term_weights(BASE_TERMS)

SEMANTIC_QUERY = "gold pretread roast preg robb"
SEMANTIC_THRESHOLD = 0.24
TIME_FILTER = (1900, 2025)

def remove_accents(input_str):
    nfkd_form = unicodedata.normalize('NFKD', input_str)
    return "".join([c for c in nfkd_form if not unicodedata.combining(c)])

def get_bib_key_from_authors_and_year(authors_str, year_str):
    default_key = "Desconocido"
    year_clean = year_str if str(year_str).isdigit() else "0000"

    if not authors_str or not authors_str.strip():
        return f"{default_key}{year_clean}"

    first_author_raw = authors_str.split(";")[0].strip()
    if not first_author_raw:
        return f"{default_key}{year_clean}"

    if "," in first_author_raw:
        last_name = first_author_raw.split(",", 1)[0].strip()
    else:
        tokens = first_author_raw.split()

        def is_initial(tok):
            tok_no_dot = tok.replace(".", "")
            return len(tok_no_dot) <= 2 and tok_no_dot == tok_no_dot.upper()

        filtered_tokens = [t for t in tokens if not is_initial(t)]

        if len(filtered_tokens) > 1:
            last_name = filtered_tokens[-1]
        elif len(filtered_tokens) == 1:
            last_name = filtered_tokens[0]
        else:
            last_name = tokens[0]

    last_name_norm = remove_accents(last_name)
    last_name_norm = re.sub(r'[^a-zA-Z0-9]+', '', last_name_norm)
    if not last_name_norm:
        last_name_norm = default_key

    last_name_norm = last_name_norm.capitalize()
    return f"{last_name_norm}{year_clean}"

# ============================================================
#   generate_bibtex: ESCRIBE TODAS LAS REFERENCIAS EN .BIB
# ============================================================
def generate_bibtex(df, output_dir):
    df = df.reset_index(drop=True)
    bib_file_path = os.path.join(output_dir, 'references.bib')

    with open(bib_file_path, 'w', encoding='utf-8') as bib_file:
        for i, row in df.iterrows():
            # Manejo de nulos/NaN en Authors para evitar el error 'float' object has no attribute 'replace'
            authors_raw = row.get('Authors', '')
            if pd.isna(authors_raw):
                authors_raw = ''
            authors_raw = str(authors_raw)

            # Estructura de autores con "and"
            authors_bib = authors_raw.replace(';', ' and ')

            year_ = str(row.get('Year', ''))
            bib_key = get_bib_key_from_authors_and_year(authors_raw, year_)
            df.at[i, 'BibKey'] = bib_key

            title   = str(row.get('Title', '')).replace('{','').replace('}','')
            journal = str(row.get('Journal', '')).replace('&', r'\&')
            volume  = str(row.get('Volume', ''))
            number  = str(row.get('Number', ''))
            pages   = str(row.get('Pages', ''))
            doi     = str(row.get('DOI', ''))

            bib_file.write(f"@article{{{bib_key},\n")
            bib_file.write(f"  author    = {{{authors_bib}}},\n")
            bib_file.write(f"  title     = {{{title}}},\n")
            bib_file.write(f"  journal   = {{{journal}}},\n")

            if pd.notna(volume) and volume.strip():
                bib_file.write(f"  volume    = {{{volume}}},\n")
            if pd.notna(number) and number.strip():
                bib_file.write(f"  number    = {{{number}}},\n")
            if pd.notna(pages) and pages.strip():
                bib_file.write(f"  pages     = {{{pages}}},\n")
            if pd.notna(year_) and year_.strip():
                bib_file.write(f"  year      = {{{year_}}},\n")
            if pd.notna(doi) and doi.strip():
                bib_file.write(f"  doi       = {{{doi}}}\n")
            else:
                bib_file.write("  doi       = {}\n")

            bib_file.write("}\n\n")

    return df

# ================================
#  Unificación SCOPUS y WoS
# ================================
def unify_scopus(scopus_file):
    df = pd.read_csv(scopus_file, encoding='utf-8')

    # Mapeo de columnas
    column_mapping_scopus = {
        "Authors": "Authors",
        "Title": "Title",
        "Year": "Year",
        "Source title": "Journal",
        "Volume": "Volume",
        "Issue": "Number",
        "Abstract": "Abstract",
        "Author Keywords": "AuthorKeywords",
        "DOI": "DOI"
    }
    df.rename(columns=column_mapping_scopus, inplace=True)

    for col_ in ["Page start", "Page end", "Art. No."]:
        if col_ not in df.columns:
            df[col_] = None

    def combine_scopus_pages(row):
        start = str(row["Page start"]) if pd.notna(row["Page start"]) else ""
        end   = str(row["Page end"])   if pd.notna(row["Page end"]) else ""
        arti  = str(row["Art. No."])  if pd.notna(row["Art. No."]) else ""
        start, end, arti = start.strip(), end.strip(), arti.strip()

        if start and end:
            return f"{start}-{end}"
        elif arti:
            return arti
        else:
            return None

    df["Pages"] = df.apply(combine_scopus_pages, axis=1)
    df["KeywordsPlus"] = ""
    df["SourceSystem"] = "Scopus"

    return df

def unify_wos(wos_file):
    df = pd.read_excel(wos_file)

    if "Publication Year" in df.columns:
        df.rename(columns={"Publication Year": "Year"}, inplace=True)
    elif "Publication Date" in df.columns:
        df.rename(columns={"Publication Date": "Year"}, inplace=True)

    for col_ in ["Start Page", "End Page", "Article Number"]:
        if col_ not in df.columns:
            df[col_] = None

    def combine_wos_pages(row):
        start = str(row["Start Page"]) if pd.notna(row["Start Page"]) else ""
        end   = str(row["End Page"])   if pd.notna(row["End Page"]) else ""
        arti  = str(row["Article Number"]) if pd.notna(row["Article Number"]) else ""
        start, end, arti = start.strip(), end.strip(), arti.strip()

        if start and end:
            return f"{start}-{end}"
        elif arti:
            return arti
        else:
            return None

    df["Pages"] = df.apply(combine_wos_pages, axis=1)

    column_mapping_wos = {
        "Article Title": "Title",
        "Abstract": "Abstract",
        "Authors": "Authors",
        "Source Title": "Journal",
        "Volume": "Volume",
        "Issue": "Number",
        "DOI": "DOI",
        "Year": "Year"
    }
    df.rename(columns=column_mapping_wos, inplace=True)

    if "Author Keywords" in df.columns:
        df.rename(columns={"Author Keywords": "AuthorKeywords"}, inplace=True)
    else:
        df["AuthorKeywords"] = ""

    if "Keywords Plus" in df.columns:
        df.rename(columns={"Keywords Plus": "KeywordsPlus"}, inplace=True)
    else:
        df["KeywordsPlus"] = ""

    df["SourceSystem"] = "WoS"

    return df

def text_score_for_weights(text, term_weights):
    text_lower = text.lower()
    score = 0.0
    for term, weight in term_weights.items():
        # Buscamos el término como palabra completa
        if re.search(r"\b" + re.escape(term) + r"\b", text_lower):
            score += weight
    return score

def abstract_pairs_score(text, term_weights):
    text_lower = text.lower()
    filler = r"(?:[,\:\;\s]*(?:and|or|in|the|a|an|of|to|by)*[,\:\;\s]*)?"
    c = 0.0
    terms_list = list(term_weights.keys())

    for i in range(len(terms_list)):
        for j in range(i+1, len(terms_list)):
            t1 = re.escape(terms_list[i])
            t2 = re.escape(terms_list[j])
            pattern1 = rf"\b{t1}\b{filler}\b{t2}\b"
            pattern2 = rf"\b{t2}\b{filler}\b{t1}\b"

            matches1 = re.findall(pattern1, text_lower)
            matches2 = re.findall(pattern2, text_lower)

            if matches1:
                wmean = (term_weights[terms_list[i]] + term_weights[terms_list[j]]) / 2
                c += wmean
            if matches2:
                wmean = (term_weights[terms_list[i]] + term_weights[terms_list[j]]) / 2
                c += wmean
    return c

def custom_score_dataframe(df, term_weights, min_score=2.0):
    df = df.copy()
    scores = []
    for idx, row in tqdm(df.iterrows(), total=len(df), desc="Scoring docs"):
        t_score = text_score_for_weights(str(row.get("Title","")), term_weights)
        joined_keywords = str(row.get("AuthorKeywords","")) + " " + str(row.get("KeywordsPlus",""))
        k_score = text_score_for_weights(joined_keywords, term_weights)
        a_score = abstract_pairs_score(str(row.get("Abstract","")), term_weights)
        total_score = t_score + k_score + a_score
        scores.append(total_score)

    df["Score"] = scores
    df = df[df["Score"] >= min_score].copy()
    df.sort_values("Score", ascending=False, inplace=True)
    df["StudyRank"] = range(1, len(df)+1)
    return df.reset_index(drop=True)

def semantic_filter(df, query_text, threshold=0.5):
    if df.empty:
        return df

    query_emb = embedding_model.encode([query_text], show_progress_bar=False)[0]
    combined_text = (
        df["Title"].fillna("") + " " +
        df["Abstract"].fillna("") + " " +
        df["AuthorKeywords"].fillna("") + " " +
        df["KeywordsPlus"].fillna("")
    ).tolist()

    doc_embs = embedding_model.encode(combined_text, show_progress_bar=False)
    cos_sims = st_utils.pytorch_cos_sim(query_emb, doc_embs)
    cos_sims = cos_sims[0].cpu().numpy()

    df["SemSim"] = cos_sims
    df2 = df[df["SemSim"] >= threshold].copy()
    df2.sort_values("SemSim", ascending=False, inplace=True)
    return df2.reset_index(drop=True)

def advanced_summary(text, max_length=80):
    MAX_CHARS_INPUT = 2000
    text = text.strip()
    if len(text) > MAX_CHARS_INPUT:
        text = text[:MAX_CHARS_INPUT]
    try:
        summary_list = summarizer(
            text,
            max_length=max_length,
            min_length=30,
            do_sample=False
        )
        return summary_list[0]["summary_text"]
    except Exception as e:
        print(f"Error during summarization: {e}")
        return text[:300] + "..." if len(text) > 300 else text

def make_apa_reference(row):
    authors = row.get('Authors', '')
    year = row.get('Year', '')
    title = row.get('Title', '')
    journal = row.get('Journal', '')
    volume = row.get('Volume', '')
    number = row.get('Number', '')
    pages = row.get('Pages', '')
    doi = row.get('DOI', '')

    ref = f"{authors} ({year}). {title}. *{journal}*, {volume}"
    if number and not pd.isna(number):
        ref += f"({number})"
    if pages and not pd.isna(pages):
        ref += f", {pages}"
    if doi and not pd.isna(doi):
        ref += f". https://doi.org/{doi}"
    return ref

def embed_texts(df, model):
    texts = (
        df["Title"].fillna("") + " " +
        df["Abstract"].fillna("") + " " +
        df["AuthorKeywords"].fillna("") + " " +
        df["KeywordsPlus"].fillna("")
    ).tolist()
    embeddings = model.encode(texts, show_progress_bar=True)
    return embeddings

def find_optimal_clusters(embeddings, min_k=2, max_k=12):
    best_score = -1
    best_k = min_k
    for k in range(min_k, max_k + 1):
        kmeans = KMeans(n_clusters=k, random_state=42)
        labels = kmeans.fit_predict(embeddings)
        score = silhouette_score(embeddings, labels)
        if score > best_score:
            best_k = k
            best_score = score
    return best_k

def generate_docs_by_cluster(df_all, output_dir):
    cluster_counts = df_all['Cluster'].value_counts().sort_values(ascending=False)
    sorted_clusters = cluster_counts.index.tolist()

    for cluster_id in sorted_clusters:
        cluster_doc = Document()
        cluster_doc.add_heading(f"Cluster {cluster_id}", level=1)
        cdf = df_all[df_all["Cluster"] == cluster_id].reset_index(drop=True)

        for i, row in tqdm(cdf.iterrows(), total=len(cdf),
                           desc=f"Summaries (Cluster {cluster_id})"):
            abstract_text = str(row.get("Abstract",""))
            summary = advanced_summary(abstract_text, max_length=80)
            cluster_doc.add_heading(f"Paper {i+1}: {row.get('Title','')}", level=2)
            cluster_doc.add_paragraph(f"**PEGASUS Summary**: {summary}", style='Normal')
            apa_ref = make_apa_reference(row)
            cluster_doc.add_paragraph(f"**APA Reference**: {apa_ref}", style='Normal')

            if "BibKey" in row:
                cluster_doc.add_paragraph(f"BibTeX Key: {row['BibKey']}", style='Normal')

        doc_name = os.path.join(output_dir, f"cluster_{cluster_id}.docx")
        cluster_doc.save(doc_name)
        print(f"Generated cluster doc: {doc_name}")

def generate_wordcloud(df, output_dir):
    text_series = (
        df["Title"].fillna("") + " " +
        df["Abstract"].fillna("") + " " +
        df["AuthorKeywords"].fillna("") + " " +
        df["KeywordsPlus"].fillna("")
    )
    all_text = " ".join(text_series.str.lower().tolist())
    if not all_text.strip():
        print("Not enough text to generate WordCloud.")
        return

    wc = WordCloud(width=1200, height=600, background_color="white", colormap="viridis")
    wc.generate(all_text)

    plt.figure(figsize=(10,5))
    plt.imshow(wc, interpolation="bilinear")
    plt.axis("off")
    plt.title("Word Cloud", fontsize=16)
    plt.tight_layout()

    out_path = os.path.join(output_dir, "wordcloud.png")
    plt.savefig(out_path, dpi=300)
    plt.close()
    print(f"Generated WordCloud: {out_path}")

def plot_cluster_themes(df, output_dir):
    cluster_counts = df['Cluster'].value_counts().sort_index()
    if cluster_counts.empty:
        print("No data to plot for cluster themes.")
        return

    plt.figure(figsize=(8,5))
    sns.barplot(x=cluster_counts.index, y=cluster_counts.values, palette="viridis")
    plt.title("Documents Per Cluster", fontsize=16)
    plt.xlabel("Cluster ID", fontsize=12)
    plt.ylabel("Count", fontsize=12)

    for i, val in enumerate(cluster_counts.values):
        plt.text(i, val+0.2, str(val), ha='center')

    plt.tight_layout()
    out_path = os.path.join(output_dir, "cluster_themes_barplot.png")
    plt.savefig(out_path, dpi=300)
    plt.close()
    print(f"Generated cluster themes barplot: {out_path}")

def generate_single_doc_summaries(df, output_dir, doc_name="summaries.doc"):
    doc = Document()
    doc.add_heading("Extracted Insights", level=1)

    sources = df["SourceSystem"].unique()
    for source in sources:
        sub = df[df["SourceSystem"] == source]
        if sub.empty:
            continue

        doc.add_heading(f"Insights from {source}", level=2)
        for i, row in sub.iterrows():
            t = str(row.get("Title",""))
            abs_ = str(row.get("Abstract",""))
            ref = make_apa_reference(row)
            doc.add_paragraph(t, style='Normal')
            doc.add_paragraph(abs_, style='Normal')
            doc.add_paragraph(f"Reference: {ref}", style='Normal')

    path_ = os.path.join(output_dir, doc_name)
    doc.save(path_)
    print(f"Generated single doc: {path_}")

def plot_study_scores(df, output_dir):
    if "Score" not in df.columns:
        return

    df_12 = df[df["StudyRank"] <= 12].copy()
    if df_12.empty:
        print("No data for top 12 articles.")
        return

    plt.figure(figsize=(12,5))
    sns.barplot(data=df_12, x="StudyRank", y="Score", palette="viridis")
    plt.title("Study Score by Rank (Top 12 Only)", fontsize=16)
    plt.xlabel("Study Rank", fontsize=12)
    plt.ylabel("Score (Weighted Terms)", fontsize=12)

    for i, row in df_12.iterrows():
        plt.text(i, row["Score"]+0.1, f"{row['Score']:.1f}", ha='center')

    plt.tight_layout()
    out_path = os.path.join(output_dir, "study_scores_barplot.png")
    plt.savefig(out_path, dpi=300)
    plt.close()
    print(f"Generated study scores barplot (Top 12): {out_path}")

def plot_num_final_articles(df, output_dir):
    count_final = len(df)
    plt.figure(figsize=(4,5))
    sns.barplot(x=["Selected Articles"], y=[count_final], palette="viridis")
    plt.ylabel("Count")
    plt.title("Final Selected Articles")

    plt.text(0, count_final+0.1, str(count_final), ha='center', fontsize=12)
    out_path = os.path.join(output_dir, "final_articles_count.png")
    plt.savefig(out_path, dpi=300)
    plt.close()
    print(f"Generated final article count barplot: {out_path}")

def plot_top_journals(df, output_dir, top_n=12):
    journals = df["Journal"].fillna("Unknown").value_counts().head(top_n)
    plt.figure(figsize=(12,7))
    sns.barplot(x=journals.values, y=journals.index, palette="viridis", orient='h')
    plt.title(f"Top {top_n} Journals - Final Selection", fontsize=16)
    plt.xlabel("Number of Publications", fontsize=12)
    plt.ylabel("Journal", fontsize=12)

    for i, val in enumerate(journals.values):
        plt.text(val+0.3, i, str(val), va='center')

    plt.tight_layout()
    out_path = os.path.join(output_dir, "top_journals.png")
    plt.savefig(out_path, dpi=300)
    plt.close()
    print(f"Generated top {top_n} journals barplot: {out_path}")

def analyze_keywords_richness(df, output_dir, top_n=20):
    author_keywords = df["AuthorKeywords"].dropna().tolist()
    author_kw_list = []
    for kw_line in author_keywords:
        splitted = re.split(r';|,', kw_line)
        splitted = [k.strip().lower() for k in splitted if k.strip()]
        author_kw_list.extend(splitted)

    author_kw_series = pd.Series(author_kw_list)
    if not author_kw_series.empty:
        author_kw_counts = author_kw_series.value_counts().head(top_n)
        plt.figure(figsize=(8,5))
        sns.barplot(x=author_kw_counts.values, y=author_kw_counts.index, palette="viridis")
        plt.title("Top Author's Keywords", fontsize=16)
        plt.xlabel("Frequency", fontsize=12)
        plt.ylabel("Keyword", fontsize=12)
        for i, val in enumerate(author_kw_counts.values):
            plt.text(val+0.1, i, str(val), va='center')
        plt.tight_layout()
        out_path = os.path.join(output_dir, "authors_keywords_top.png")
        plt.savefig(out_path, dpi=300)
        plt.close()
        print(f"Generated AuthorKeywords barplot: {out_path}")
    else:
        print("No AuthorKeywords to plot.")

    plus_keywords = df["KeywordsPlus"].dropna().tolist()
    plus_kw_list = []
    for kw_line in plus_keywords:
        splitted = re.split(r';|,', kw_line)
        splitted = [k.strip().lower() for k in splitted if k.strip()]
        plus_kw_list.extend(splitted)

    plus_kw_series = pd.Series(plus_kw_list)
    if not plus_kw_series.empty:
        plus_kw_counts = plus_kw_series.value_counts().head(top_n)
        plt.figure(figsize=(8,5))
        sns.barplot(x=plus_kw_counts.values, y=plus_kw_counts.index, palette="viridis")
        plt.title("Top Keywords Plus", fontsize=16)
        plt.xlabel("Frequency", fontsize=12)
        plt.ylabel("Keyword", fontsize=12)
        for i, val in enumerate(plus_kw_counts.values):
            plt.text(val+0.1, i, str(val), va='center')
        plt.tight_layout()
        out_path = os.path.join(output_dir, "keywords_plus_top.png")
        plt.savefig(out_path, dpi=300)
        plt.close()
        print(f"Generated KeywordsPlus barplot: {out_path}")
    else:
        print("No KeywordsPlus to plot.")

def plot_year_distribution(df, output_dir):
    df["Year"] = pd.to_numeric(df["Year"], errors='coerce')
    year_counts = df["Year"].dropna().value_counts().sort_index()
    if year_counts.empty:
        print("No year data to plot.")
        return

    plt.figure(figsize=(8,5))
    sns.barplot(x=year_counts.index.astype(int), y=year_counts.values, palette="viridis")
    plt.title("Year Distribution - Selected Articles", fontsize=16)
    plt.xlabel("Year", fontsize=12)
    plt.ylabel("Count", fontsize=12)

    for i, val in enumerate(year_counts.values):
        plt.text(i, val+0.1, str(val), ha='center')
    plt.tight_layout()
    out_path = os.path.join(output_dir, "year_distribution.png")
    plt.savefig(out_path, dpi=300)
    plt.close()
    print(f"Generated year distribution barplot: {out_path}")

def build_coauthor_network(df, output_dir):
    G = nx.Graph()
    for idx, row_ in df.iterrows():
        authors_str = row_.get("Authors", "")
        authors_list = [a.strip() for a in authors_str.split(";") if a.strip()]
        for author in authors_list:
            if not G.has_node(author):
                G.add_node(author)
        for i in range(len(authors_list)):
            for j in range(i+1, len(authors_list)):
                G.add_edge(authors_list[i], authors_list[j])

    degrees = G.degree()
    top_authors = sorted(degrees, key=lambda x: x[1], reverse=True)[:12]
    top_nodes = [a[0] for a in top_authors]

    G_sub = G.subgraph(top_nodes).copy()

    plt.figure(figsize=(10, 10))
    pos = nx.circular_layout(G_sub)
    nx.draw(
        G_sub, pos,
        with_labels=True,
        node_size=600,
        node_color='skyblue',
        edge_color='gray',
        font_size=7,
        alpha=0.8
    )
    plt.title("Co-author Network (Top 12 Authors by Degree)", fontsize=14)
    plt.axis("off")

    out_path_nx = os.path.join(output_dir, "coauthor_network_networkx.png")
    plt.savefig(out_path_nx, dpi=300)
    plt.close()
    print(f"Generated co-author network with networkx: {out_path_nx}")

    rows = []
    for idx, row_ in df.iterrows():
        y = row_.get("Year", None)
        authors_str = str(row_.get("Authors",""))
        authors_list = [a.strip() for a in authors_str.split(";") if a.strip()]
        for auth in authors_list:
            rows.append((auth, y))

    bubble_df = pd.DataFrame(rows, columns=["Author","Year"])
    bubble_df["Year"] = pd.to_numeric(bubble_df["Year"], errors='coerce')
    group_df = bubble_df.dropna(subset=["Author","Year"])
    group_df = group_df.groupby(["Author","Year"], as_index=False).size()

    author_counts = group_df.groupby("Author")["size"].sum().sort_values(ascending=False)
    top_12_authors = author_counts.head(12).index.tolist()
    group_top = group_df[group_df["Author"].isin(top_12_authors)].copy()

    fig = px.scatter(
        group_top,
        x="Year",
        y="Author",
        size="size",
        color="size",
        hover_name="Author",
        title="Author Contributions Over Time",
        color_continuous_scale="Viridis",
    )
    fig.update_traces(marker=dict(symbol='circle', sizemode='area',
                                  line=dict(width=1, color='white')))
    fig.update_layout(
        xaxis_title="Year",
        yaxis_title="Author",
        legend_title="Documents"
    )

    out_path = os.path.join(output_dir, "coauthor_network.png")
    fig.write_image(out_path, format="png", scale=2)
    print(f"Generated bubble chart: {out_path}")

def plot_abstract_keywords_frequency(df, output_dir, top_n=20):
    stop_words = set(stopwords.words('english'))
    all_abstracts = " ".join(df["Abstract"].dropna().tolist()).lower()
    tokens = word_tokenize(all_abstracts)
    tokens_clean = [t for t in tokens if t.isalpha() and t not in stop_words]

    freq_series = pd.Series(tokens_clean).value_counts().head(top_n)
    if freq_series.empty:
        print("Not enough tokens to plot abstract keywords frequency.")
        return

    plt.figure(figsize=(8,5))
    sns.barplot(x=freq_series.values, y=freq_series.index, palette="viridis")
    plt.title("Most Frequent Words in Abstracts", fontsize=16)
    plt.xlabel("Frequency", fontsize=12)
    plt.ylabel("Word", fontsize=12)

    for i, val in enumerate(freq_series.values):
        plt.text(val+0.1, i, str(val), va='center')
    plt.tight_layout()
    out_path = os.path.join(output_dir, "abstract_keywords_frequency.png")
    plt.savefig(out_path, dpi=300)
    plt.close()
    print(f"Generated abstract keywords frequency barplot: {out_path}")

# =========================
#         MAIN
# =========================
def main():
    print(">>> Starting process with advanced LLM techniques ...")

    # 1) Unificar datos de Scopus y WoS
    scopus_df = unify_scopus(SCOPUS_PATH)
    wos_df    = unify_wos(WOS_PATH)

    print("\n[INFO] Columnas SCOPUS unificadas:", scopus_df.columns.tolist())
    print("[INFO] Registros SCOPUS:", len(scopus_df))
    print("[INFO] Columnas WoS unificadas:", wos_df.columns.tolist())
    print("[INFO] Registros WoS:", len(wos_df))

    # 2) Combinar ambos DataFrames (sin drop_duplicates para no perder referencias)
    combined_df = pd.concat([scopus_df, wos_df], ignore_index=True)
    print("\n[INFO] Registros totales combinados:", len(combined_df))

    # 3) Asegurarnos de que existe Título y Abstract
    combined_df.dropna(subset=["Title", "Abstract"], inplace=True)
    print("[INFO] Registros tras eliminar sin Título/Abstract:", len(combined_df))

    # 4) Filtro temporal
    combined_df["Year"] = pd.to_numeric(combined_df["Year"], errors='coerce')
    start_year, end_year = TIME_FILTER
    combined_df = combined_df[
        (combined_df["Year"] >= start_year) & 
        (combined_df["Year"] <= end_year)
    ]
    print(f"[INFO] Registros tras filtrar años entre {start_year} y {end_year}: {len(combined_df)}")

    # 5) Filtro semántico
    print(f"\nApplying semantic filter: '{SEMANTIC_QUERY}' (threshold={SEMANTIC_THRESHOLD})")
    df_sem = semantic_filter(combined_df, SEMANTIC_QUERY, threshold=SEMANTIC_THRESHOLD)
    print(f"Docs after semantic filter: {len(df_sem)} / {len(combined_df)}")

    # 6) Filtro de score
    print("Scoring docs with weighted terms + pairs in Abstract ... (min_score=2)")
    df_scored = custom_score_dataframe(df_sem, TERM_WEIGHTS, min_score=3.0)
    print(f"Docs after scoring >=2.0: {len(df_scored)}")

    if df_scored.empty:
        print("No docs after final scoring. Exiting.")
        return

    # 7) Generar BibTeX
    print("Generating BibTeX for all scored docs, format @ApellidoAño ...")
    df_scored = generate_bibtex(df_scored, OUTPUT_DIR)

    # 8) Tomamos todos los docs resultantes
    df_top = df_scored.copy()
    print(f"Working with all {len(df_top)} articles that meet the filters.\n")

    print("Top docs from df_top:")
    print(df_top[["Title", "Score", "StudyRank", "BibKey", "SourceSystem"]].head(10))

    # 9) Embeddings y clustering
    print("Generating embeddings for the selected docs ...")
    embeddings = embed_texts(df_top, embedding_model)
    n_docs = len(df_top)

    if n_docs == 1:
        print("Only 1 doc => cluster=1")
        df_top["Cluster"] = 1
    elif n_docs == 2:
        print("2 docs => forcing K=2")
        kmeans = KMeans(n_clusters=2, random_state=42)
        labels = kmeans.fit_predict(embeddings)
        df_top["Cluster"] = labels + 1
    else:
        max_k = min(12, n_docs - 1)
        min_k = 2
        print(f"Finding optimal k in range [{min_k},{max_k}]")
        best_k = find_optimal_clusters(embeddings, min_k=min_k, max_k=max_k)
        print(f"Optimal k = {best_k}")
        kmeans = KMeans(n_clusters=best_k, random_state=42)
        labels = kmeans.fit_predict(embeddings)
        df_top["Cluster"] = labels + 1

    # 10) Generar .docx por cluster
    generate_docs_by_cluster(df_top, OUTPUT_DIR)

    # 11) Word Cloud
    generate_wordcloud(df_top, OUTPUT_DIR)

    # 12) Barplot de Clusters
    plot_cluster_themes(df_top, OUTPUT_DIR)

    # 13) Summaries.doc
    generate_single_doc_summaries(df_top, OUTPUT_DIR, doc_name="summaries.doc")

    # 14) Barplot Score (solo top 12)
    plot_study_scores(df_top, OUTPUT_DIR)

    # =========================================================
    # NUEVAS ESTADÍSTICAS BIBLIOMÉTRICAS
    # =========================================================
    plot_num_final_articles(df_top, OUTPUT_DIR)
    plot_top_journals(df_top, OUTPUT_DIR, top_n=12)
    analyze_keywords_richness(df_top, OUTPUT_DIR, top_n=20)
    plot_year_distribution(df_top, OUTPUT_DIR)
    build_coauthor_network(df_top, OUTPUT_DIR)
    plot_abstract_keywords_frequency(df_top, OUTPUT_DIR, top_n=20)

    print(">>> Process completed successfully. Outputs in:", OUTPUT_DIR)

if __name__ == "__main__":
    main()
