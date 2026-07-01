#!/usr/bin/env python3
"""Compila manuscript.docx con estructura PRISMA 2020 + marcadores Zotero RTF/ODF Scan.

La prosa larga (Introduction/Discussion/Conclusions) se deja como placeholders para
que el usuario la complete (o la skill la pase por humanizer antes de inyectarla).

Las citas se insertan como `{|Apellido Año|}` — Zotero RTF/ODF Scan las convierte
en citas vivas vinculadas a la colección "SR — <slug>".
"""
from __future__ import annotations

import argparse
import json
import sys
from datetime import datetime
from pathlib import Path

import yaml
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_h1(doc: Document, text: str):
    doc.add_heading(text, level=1)


def add_h2(doc: Document, text: str):
    doc.add_heading(text, level=2)


def add_p(doc: Document, text: str, italic: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    p.paragraph_format.space_after = Pt(6)


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("project_dir")
    args = ap.parse_args()

    proj = Path(args.project_dir)
    cfg = yaml.safe_load((proj / "project_config.yaml").read_text())
    corpus = []
    cj = proj / "master_corpus.json"
    if cj.exists():
        corpus = json.loads(cj.read_text())

    title = (cfg.get("manuscript") or {}).get("title") or (cfg.get("project") or {}).get("name") or "Untitled systematic review"
    lang = (cfg.get("manuscript") or {}).get("language") or "es"

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    # Title page
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(title)
    r.bold = True
    r.font.size = Pt(16)

    doc.add_paragraph().add_run("Authors: [a rellenar]").italic = True
    doc.add_paragraph().add_run(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d')}").italic = True
    doc.add_page_break()

    # Abstract
    add_h1(doc, "Abstract")
    add_h2(doc, "Background")
    add_p(doc, "[Rellenar — contexto y rationale, 2–3 frases]")
    add_h2(doc, "Objectives")
    add_p(doc, (cfg.get("review") or {}).get("research_question") or "[a rellenar]")
    add_h2(doc, "Data Sources")
    dbs = [k for k, v in (cfg.get("databases") or {}).items() if v]
    add_p(doc, ", ".join(dbs))
    add_h2(doc, "Study Eligibility Criteria")
    add_p(doc, "[Rellenar — inclusión/exclusión]")
    add_h2(doc, "Participants and Interventions")
    pico = (cfg.get("review") or {}).get("pico") or {}
    add_p(doc, f"Population: {pico.get('population','[?]')}. Intervention: {pico.get('intervention','[?]')}. Comparator: {pico.get('comparator','[?]')}. Outcome: {pico.get('outcome','[?]')}.")
    add_h2(doc, "Study Appraisal and Synthesis Methods")
    add_p(doc, f"Risk of bias: {cfg.get('rob_tool','[a rellenar]')}. Synthesis: {'meta-analysis with random effects' if (cfg.get('meta_analysis') or {}).get('enabled') else 'narrative synthesis'}.")
    add_h2(doc, "Results")
    add_p(doc, "[A rellenar tras Fase 13]")
    add_h2(doc, "Limitations")
    add_p(doc, "[A rellenar]")
    add_h2(doc, "Conclusions")
    add_p(doc, "[A rellenar]")
    add_h2(doc, "Registration")
    add_p(doc, "PROSPERO: [añadir CRD tras registro]")

    doc.add_page_break()

    # 1. Introduction
    add_h1(doc, "1. Introduction")
    add_h2(doc, "1.1 Rationale [Item 3]")
    add_p(doc, "[Prosa — contexto, gap, justificación. Pasar por skill `humanizer` antes de inyectar. "
               "Incluye citas como {|Author Year|} — Zotero RTF/ODF Scan las convertirá.]")
    add_h2(doc, "1.2 Objectives [Item 4]")
    add_p(doc, (cfg.get("review") or {}).get("research_question") or "[Pregunta de investigación PICO]")

    # 2. Methods
    add_h1(doc, "2. Methods")
    add_h2(doc, "2.1 Protocol and registration [Item 24]")
    add_p(doc, "El protocolo se registró en PROSPERO (CRD: [añadir tras registro]). Esta RS sigue las guías PRISMA 2020 (Page et al., 2021).")
    add_h2(doc, "2.2 Eligibility criteria [Item 5]")
    add_p(doc, "[Tabla de inclusión/exclusión con PICO completo]")
    add_h2(doc, "2.3 Information sources [Item 6]")
    add_p(doc, f"Se buscó en: {', '.join(dbs)}. Fecha de última búsqueda: [a rellenar].")
    add_h2(doc, "2.4 Search strategy [Item 7]")
    add_p(doc, "Estrategias completas por base disponibles en searches/*_query.txt. Sintaxis PubMed:")
    pm = (proj / "searches" / "pubmed_query.txt")
    if pm.exists():
        doc.add_paragraph(pm.read_text(encoding="utf-8")[:1000]).style = "Intense Quote"
    add_h2(doc, "2.5 Selection process [Item 8]")
    add_p(doc, "Pass 1: cribado automático por reglas keyword (asistido por IA, ver checklist PRISMA-trAIce). "
               "Pass 2: revisión asistida por IA de title/abstract con confirmación humana. "
               "Discrepancias resueltas por consenso. Herramienta: skill systematic-review v1.0.")
    add_h2(doc, "2.6 Data collection process [Item 9]")
    add_p(doc, "Extracción dual a master_corpus.xlsx. Discrepancias por consenso.")
    add_h2(doc, "2.7 Data items [Items 10a, 10b]")
    add_p(doc, f"Outcomes: {pico.get('outcome','[a rellenar]')}. Variables: design, sample size, intervention, comparator, follow-up, outcome measure.")
    add_h2(doc, "2.8 Study risk of bias assessment [Item 11]")
    add_p(doc, f"Tool: {cfg.get('rob_tool') or '[a elegir según design predominante]'}. Evaluación dual.")
    add_h2(doc, "2.9 Effect measures [Item 12]")
    add_p(doc, f"Effect measure: {(cfg.get('meta_analysis') or {}).get('effect_measure') or 'N/A — síntesis narrativa'}.")
    add_h2(doc, "2.10 Synthesis methods [Items 13a–f]")
    add_p(doc, "Síntesis narrativa por temas. " + ("Meta-análisis con modelo random effects, heterogeneidad evaluada con I². " if (cfg.get('meta_analysis') or {}).get('enabled') else "No procede meta-análisis."))
    add_h2(doc, "2.11 Reporting bias assessment [Item 14]")
    add_p(doc, "Funnel plot inspeccionado si MA con ≥10 estudios.")
    add_h2(doc, "2.12 Certainty assessment [Item 15]")
    add_p(doc, "GRADE para outcomes principales.")

    # 3. Results
    add_h1(doc, "3. Results")
    add_h2(doc, "3.1 Study selection [Items 16a, 16b]")
    add_p(doc, "Ver Figura 1 (PRISMA 2020 flow diagram).")
    add_p(doc, "[Insertar prisma_flow.png aquí]")
    add_h2(doc, "3.2 Study characteristics [Item 17]")
    add_p(doc, "[Tabla 1: características de los estudios incluidos]")
    add_h2(doc, "3.3 Risk of bias in studies [Item 18]")
    add_p(doc, "[Figura 2: RoB summary plot]")
    add_h2(doc, "3.4 Results of individual studies [Item 19]")
    add_p(doc, "[Resúmenes por estudio]")
    add_h2(doc, "3.5 Results of syntheses [Items 20a–d]")
    add_p(doc, "[Síntesis por tema / forest plot si MA]")
    add_h2(doc, "3.6 Reporting biases [Item 21]")
    add_p(doc, "[Funnel plot si procede]")
    add_h2(doc, "3.7 Certainty of evidence [Item 22]")
    add_p(doc, "[Tabla SoF — Summary of Findings]")

    # 4. Discussion
    add_h1(doc, "4. Discussion")
    add_h2(doc, "4.1 Summary of evidence [Item 23a]")
    add_p(doc, "[Prosa — Pasar por humanizer. Citas como {|Author Year|}]")
    add_h2(doc, "4.2 Limitations of the evidence [Item 23b]")
    add_p(doc, "[A rellenar]")
    add_h2(doc, "4.3 Limitations of the review process [Item 23c]")
    add_p(doc, "Posible limitación: cribado Pass 1 automatizado por keyword puede haber excluido estudios con terminología no convencional. Mitigado por: cribado Pass 2 dual con confirmación humana y búsqueda en 4–6 bases complementarias.")
    add_h2(doc, "4.4 Implications [Item 23d]")
    add_p(doc, "[Implicaciones para práctica, política, investigación futura]")

    # 5. Conclusions
    add_h1(doc, "5. Conclusions")
    add_p(doc, "[Una frase resumen]")

    # Declarations
    add_h1(doc, "Declarations")
    add_h2(doc, "Funding [Item 25]")
    add_p(doc, "[a rellenar]")
    add_h2(doc, "Competing interests [Item 26]")
    add_p(doc, "None declared / [a rellenar]")
    add_h2(doc, "Data availability [Item 27]")
    add_p(doc, "Master corpus, search strategies, screening logs y PRISMA flow disponibles a petición.")
    add_h2(doc, "AI tools used")
    add_p(doc, "Esta RS empleó la skill `systematic-review v1.0` (Claude Code, Anthropic) para cribado Pass 1 y revisión asistida de abstracts. Ver checklist PRISMA-trAIce adjunto.")

    # References placeholder
    add_h1(doc, "References")
    add_p(doc, "Esta sección la rellena Zotero al correr RTF/ODF Scan sobre los marcadores {|Author Year|} del manuscrito. "
               "Colección Zotero: \"SR — " + ((cfg.get('project') or {}).get('slug') or 'untitled') + "\".")

    out = proj / "manuscript" / "manuscript.docx"
    out.parent.mkdir(exist_ok=True)
    doc.save(out)

    print(json.dumps({"ok": True, "path": str(out), "included_corpus_n": sum(1 for r in corpus if r.get('included_final') is True)}))
    return 0


if __name__ == "__main__":
    sys.exit(main())
